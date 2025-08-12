#!/usr/bin/env python3
"""
Magnus Client Intake Form Generator
- Save drafts in Word format (.docx) for editing
- Generate final reports in PDF format

INSTALLATION REQUIRED:
pip install reportlab python-docx
"""

import os
import sys
import traceback

# Check for required packages
try:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib.units import inch
except ImportError:
    print("ERROR: ReportLab is not installed. Please run: pip install reportlab")
    sys.exit(1)

try:
    from docx import Document
    from docx.shared import Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("ERROR: python-docx is not installed. Please run: pip install python-docx")
    sys.exit(1)

from .validation import parse_usd, format_usd, parse_percent, format_percent


def fmt_usd(val):
    num = parse_usd(val)
    return format_usd(num) if num is not None else "[Not provided]"


def fmt_percent(val):
    num = parse_percent(val)
    return format_percent(num) if num is not None else "[Not provided]"

def save_draft_word(form_data, output_path):
    """Save form data as a Word document draft"""
    try:
        # Create document
        doc = Document()
        
        # Add title
        doc.add_heading('Magnus Client Intake Form', 0)
        doc.add_paragraph()
        
        # Personal Information
        doc.add_heading('Personal Information', level=1)
        doc.add_paragraph(f"Full Name: {form_data.get('full_name', '[Not provided]')}")
        doc.add_paragraph(f"Date of Birth: {form_data.get('dob', '[Not provided]')}")
        doc.add_paragraph(f"Social Security Number: {form_data.get('ssn', '[Not provided]')}")
        doc.add_paragraph(f"Citizenship: {form_data.get('citizenship', '[Not provided]')}")
        doc.add_paragraph(f"Marital Status: {form_data.get('marital_status', '[Not provided]')}")
        doc.add_paragraph()
        
        # Contact Information
        doc.add_heading('Contact Information', level=1)
        doc.add_paragraph(f"Residential Address: {form_data.get('address', '[Not provided]')}")
        doc.add_paragraph(f"Email: {form_data.get('email', '[Not provided]')}")
        doc.add_paragraph(f"Home Phone: {form_data.get('phone_home', '[Not provided]')}")
        doc.add_paragraph(f"Mobile Phone: {form_data.get('phone_mobile', '[Not provided]')}")
        doc.add_paragraph(f"Work Phone: {form_data.get('phone_work', '[Not provided]')}")
        doc.add_paragraph()
        
        # Employment Information
        doc.add_heading('Employment Information', level=1)
        doc.add_paragraph(f"Employment Status: {form_data.get('employment_status', '[Not provided]')}")
        doc.add_paragraph(f"Employer Name: {form_data.get('employer_name', '[Not provided]')}")
        doc.add_paragraph(f"Occupation/Title: {form_data.get('job_title', '[Not provided]')}")
        doc.add_paragraph(f"Years Employed: {form_data.get('years_with_employer', '[Not provided]')}")
        doc.add_paragraph(f"Annual Income: {fmt_usd(form_data.get('annual_income'))}")
        doc.add_paragraph(f"Employer Address: {form_data.get('employer_address', '[Not provided]')}")
        
        # Retirement Information (if applicable)
        if form_data.get("employment_status") == "Retired":
            doc.add_paragraph()
            doc.add_heading('Retirement Information', level=1)
            doc.add_paragraph(f"Former Employer: {form_data.get('former_employer', '[Not provided]')}")
            doc.add_paragraph(f"Source of Income: {form_data.get('income_source', '[Not provided]')}")
        
        doc.add_paragraph()
        
        # Financial Information
        doc.add_heading('Financial Information', level=1)
        doc.add_paragraph(f"Education Status: {form_data.get('education_status', '[Not provided]')}")
        doc.add_paragraph(f"Estimated Tax Bracket: {fmt_percent(form_data.get('est_tax_bracket'))}")
        doc.add_paragraph(f"Investment Risk Tolerance: {form_data.get('risk_tolerance', '[Not provided]')}")
        doc.add_paragraph(f"Investment Objectives: {form_data.get('investment_objectives', '[Not provided]')}")
        doc.add_paragraph(f"Net Worth (excluding primary home): {fmt_usd(form_data.get('est_net_worth'))}")
        doc.add_paragraph(f"Liquid Net Worth: {fmt_usd(form_data.get('est_liquid_net_worth'))}")
        doc.add_paragraph(f"Assets Held Away: {fmt_usd(form_data.get('assets_held_away'))}")
        doc.add_paragraph()
        
        # Spouse Information (if applicable)
        if not form_data.get("spouse_applicable", False):
            doc.add_heading('Spouse/Partner Information', level=1)
            doc.add_paragraph(f"Spouse Full Name: {form_data.get('spouse_full_name', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Date of Birth: {form_data.get('spouse_dob', '[Not provided]')}")
            doc.add_paragraph(f"Spouse SSN: {form_data.get('spouse_ssn', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Employment Status: {form_data.get('spouse_employment_status', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Employer Name: {form_data.get('spouse_employer_name', '[Not provided]')}")
            doc.add_paragraph(f"Spouse Occupation/Title: {form_data.get('spouse_job_title', '[Not provided]')}")
            doc.add_paragraph()
        else:
            doc.add_heading('Spouse/Partner Information', level=1)
            doc.add_paragraph("[Not applicable]")
            doc.add_paragraph()
        
        # Dependents Information
        doc.add_heading('Dependents Information', level=1)
        dependents = form_data.get("dependents", [])
        if dependents:
            for i, dep in enumerate(dependents):
                doc.add_paragraph(f"Dependent {i+1}:")
                doc.add_paragraph(f"  Name: {dep.get('full_name', '[Not provided]')}")
                doc.add_paragraph(f"  Date of Birth: {dep.get('dob', '[Not provided]')}")
                doc.add_paragraph(f"  Relationship: {dep.get('relationship', '[Not provided]')}")
        else:
            doc.add_paragraph("No dependents specified")
        doc.add_paragraph()
        
        # Beneficiaries Information
        doc.add_heading('Beneficiaries Information', level=1)
        beneficiaries = form_data.get("beneficiaries", [])
        if beneficiaries:
            for i, ben in enumerate(beneficiaries):
                doc.add_paragraph(f"Beneficiary {i+1}:")
                doc.add_paragraph(f"  Name: {ben.get('full_name', '[Not provided]')}")
                doc.add_paragraph(f"  Date of Birth: {ben.get('dob', '[Not provided]')}")
                doc.add_paragraph(f"  Relationship: {ben.get('relationship', '[Not provided]')}")
                doc.add_paragraph(f"  Percentage: {fmt_percent(ben.get('allocation'))}")
        else:
            doc.add_paragraph("No beneficiaries specified")
        doc.add_paragraph()
        
        # Asset Breakdown
        doc.add_heading('Asset Breakdown', level=1)
        asset_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures", "Short-Term", "Other"]
        for asset_type in asset_types:
            field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
            value = form_data.get(field_name)
            doc.add_paragraph(f"{asset_type}: {value}%" if value else f"{asset_type}: [Not provided]")
        doc.add_paragraph()
        
        # Investment Experience
        doc.add_heading('Investment Experience', level=1)
        asset_map = [
            ("Stocks", "stocks"),
            ("Bonds", "bonds"),
            ("Mutual Funds", "mutual_funds"),
            ("UITs", "uits"),
            ("Annuities (Fixed)", "annuities_fixed"),
            ("Annuities (Variable)", "annuities_variable"),
            ("Options", "options"),
            ("Commodities", "commodities"),
            ("Alternative Investments", "alternative_investments"),
            ("Limited Partnerships", "limited_partnerships"),
            ("Variable Contracts", "variable_contracts"),
        ]
        for label, key in asset_map:
            year = form_data.get(f"{key}_year_started")
            level = form_data.get(f"{key}_level")
            doc.add_paragraph(f"{label} – Year Started: {year if year else '[Not provided]'}, Level: {level if level else '[Not provided]'}")
            doc.add_paragraph()
        
        # Outside Broker Information
        if form_data.get("has_outside_broker", False):
            doc.add_heading('Outside Broker Information', level=1)
            doc.add_paragraph(f"Broker Firm Name: {form_data.get('outside_broker_name', '[Not provided]')}")
            doc.add_paragraph(f"Account Number: {form_data.get('outside_broker_account', '[Not provided]')}")
            doc.add_paragraph(f"Account Type: {form_data.get('outside_broker_account_type', '[Not provided]')}")
            doc.add_paragraph()
        
        # Trusted Contact Information
        doc.add_heading('Trusted Contact Information', level=1)
        doc.add_paragraph(f"Full Name: {form_data.get('trusted_full_name', '[Not provided]')}")
        doc.add_paragraph(f"Relationship: {form_data.get('trusted_relationship', '[Not provided]')}")
        doc.add_paragraph(f"Phone Number: {form_data.get('trusted_phone', '[Not provided]')}")
        doc.add_paragraph(f"Email Address: {form_data.get('trusted_email', '[Not provided]')}")
        doc.add_paragraph()
        
        # Regulatory Consent
        doc.add_heading('Regulatory Consent', level=1)
        electronic_consent = form_data.get("ed_consent", "No") or "No"
        doc.add_paragraph(f"Electronic Delivery Consent: {electronic_consent}")
        
        # Save document
        doc.save(output_path)
        return True
        
    except Exception as e:
        print(f"Error saving Word document: {str(e)}")
        traceback.print_exc()
        return False

def generate_pdf_report(form_data, output_path):
    """Generate a PDF report from form data"""
    try:
        # Validate input data
        if not isinstance(form_data, dict):
            raise ValueError("Form data must be a dictionary")
        
        # Create the PDF document
        doc = SimpleDocTemplate(
            output_path,
            pagesize=letter,
            rightMargin=72,
            leftMargin=72,
            topMargin=72,
            bottomMargin=72
        )
        
        # Define styles
        styles = getSampleStyleSheet()
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=30
        )
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            spaceAfter=12
        )
        normal_style = styles['Normal']

        # Helper function to format monetary values
        def format_money(value):
            return fmt_usd(value)

        # Helper function to format percentages
        def format_percentage(value):
            return fmt_percent(value)

        def safe(s):
            return "" if s is None else str(s)

        def format_pct(v):
            try:
                return f"{float(v):.2f}%"
            except Exception:
                return safe(v)

        # Normalize and derive fields from application state
        investment_purposes = []
        if form_data.get("inv_purpose_income"):
            investment_purposes.append("Income")
        if form_data.get("inv_purpose_growth_income"):
            investment_purposes.append("Growth and Income")
        if form_data.get("inv_purpose_cap_app"):
            investment_purposes.append("Capital Appreciation")
        if form_data.get("inv_purpose_speculation"):
            investment_purposes.append("Speculation")
        investment_purpose = ", ".join(investment_purposes)

        objectives_map = [
            ("rank_trading_profits", "Trading Profits"),
            ("rank_speculation", "Speculation"),
            ("rank_capital_appreciation", "Capital Appreciation"),
            ("rank_income", "Income"),
            ("rank_preservation", "Preservation of Capital"),
        ]
        investment_objective = "\n".join(
            f"{label}: {form_data.get(key)}" for key, label in objectives_map if form_data.get(key)
        )

        dependents = form_data.get("dependents") or []
        if not dependents and any(
            form_data.get(k) for k in ("dep_full_name", "dep_dob", "dep_relationship")
        ):
            dependents = [
                {
                    "name": form_data.get("dep_full_name"),
                    "dob": form_data.get("dep_dob"),
                    "relationship": form_data.get("dep_relationship"),
                }
            ]

        beneficiaries = form_data.get("beneficiaries") or []
        if not beneficiaries and any(
            form_data.get(k)
            for k in ("ben_full_name", "ben_dob", "ben_relationship", "ben_allocation_pct")
        ):
            beneficiaries = [
                {
                    "name": form_data.get("ben_full_name"),
                    "dob": form_data.get("ben_dob"),
                    "relationship": form_data.get("ben_relationship"),
                    "allocation": form_data.get("ben_allocation_pct"),
                }
            ]

        # Trusted contact uses tcp_* keys in state
        form_data.setdefault("trusted_full_name", form_data.get("tcp_full_name"))
        form_data.setdefault("trusted_relationship", form_data.get("tcp_relationship"))
        form_data.setdefault("trusted_phone", form_data.get("tcp_phone"))
        form_data.setdefault("trusted_email", form_data.get("tcp_email"))

        # Start building the content
        content = []

        def render_section(title, rows):
            content.append(Paragraph(title, heading_style))
            if not rows:
                content.append(Paragraph("No disclosures.", normal_style))
            else:
                for label, val in rows:
                    content.append(Paragraph(f"{label}: {safe(val)}", normal_style))
            content.append(Spacer(1, 12))
        
        # Title
        content.append(Paragraph("Magnus Client Intake Form", title_style))
        content.append(Spacer(1, 12))
        
        # Personal Information
        content.append(Paragraph("Personal Information", heading_style))
        content.append(Paragraph(f"Full Name: {form_data.get('full_name', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Date of Birth: {form_data.get('dob', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Social Security Number: {form_data.get('ssn', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Citizenship: {form_data.get('citizenship_status', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Marital Status: {form_data.get('marital_status', '[Not provided]')}", normal_style))
        content.append(Spacer(1, 12))
        
        # Contact Information
        content.append(Paragraph("Contact Information", heading_style))
        content.append(Paragraph(f"Residential Address: {form_data.get('address', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Email: {form_data.get('email', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Home Phone: {form_data.get('phone_home', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Mobile Phone: {form_data.get('phone_mobile', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Work Phone: {form_data.get('phone_work', '[Not provided]')}", normal_style))
        content.append(Spacer(1, 12))
        
        # Employment Information
        content.append(Paragraph("Employment Information", heading_style))
        content.append(Paragraph(f"Employment Status: {form_data.get('employment_status', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Employer Name: {form_data.get('employer_name', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Occupation: {form_data.get('job_title', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Years Employed: {form_data.get('years_with_employer', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Annual Income: {format_money(form_data.get('annual_income'))}", normal_style))
        content.append(Spacer(1, 12))

        # Financial Information
        content.append(Paragraph("Financial Information", heading_style))
        content.append(Paragraph(f"Education Status: {form_data.get('education_status', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Estimated Tax Bracket: {format_percentage(form_data.get('est_tax_bracket'))}", normal_style))
        content.append(Paragraph(f"Investment Risk Tolerance: {form_data.get('risk_tolerance', '[Not provided]')}", normal_style))
        
        # Investment Purpose
        if investment_purpose:
            content.append(Paragraph("Investment Purpose:", normal_style))
            for purpose in investment_purpose.split(','):
                content.append(Paragraph(f"• {purpose.strip()}", normal_style))
        else:
            content.append(Paragraph("Investment Purpose: [Not provided]", normal_style))
        
        # Investment Objectives
        if investment_objective:
            content.append(Paragraph("Investment Objectives:", normal_style))
            for objective in investment_objective.split('\n'):
                content.append(Paragraph(f"• {objective}", normal_style))
        else:
            content.append(Paragraph("Investment Objectives: [Not provided]", normal_style))
        
        content.append(Paragraph(f"Net Worth: {format_money(form_data.get('est_net_worth'))}", normal_style))
        content.append(Paragraph(f"Liquid Net Worth: {format_money(form_data.get('est_liquid_net_worth'))}", normal_style))
        content.append(Paragraph(f"Assets Held Away: {format_money(form_data.get('assets_held_away'))}", normal_style))
        content.append(Spacer(1, 12))

        # Spouse Information
        if not form_data.get('no_spouse'):
            content.append(Paragraph("Spouse Information", heading_style))
            content.append(Paragraph(f"Full Name: {form_data.get('spouse_full_name', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Date of Birth: {form_data.get('spouse_dob', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Social Security Number: {form_data.get('spouse_ssn', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Employment Status: {form_data.get('spouse_employment_status', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Employer Name: {form_data.get('spouse_employer_name', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Occupation: {form_data.get('spouse_job_title', '[Not provided]')}", normal_style))
            content.append(Spacer(1, 12))

        # Dependents
        content.append(Paragraph("Dependents", heading_style))
        if dependents:
            table_data = [["Name", "Date of Birth", "Relationship"]]
            for dep in dependents:
                table_data.append([
                    dep.get('full_name', 'Not provided') or 'Not provided',
                    dep.get('dob', 'Not provided') or 'Not provided',
                    dep.get('relationship', 'Not provided') or 'Not provided',
                ])
            table = Table(table_data, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            content.append(table)
        else:
            content.append(Paragraph("[No dependents specified]", normal_style))
        content.append(Spacer(1, 12))

        # Beneficiaries
        content.append(Paragraph("Beneficiaries", heading_style))
        if beneficiaries:
            table_data = [["Name", "Date of Birth", "Relationship", "Allocation (%)"]]
            for ben in beneficiaries:
                table_data.append([
                    ben.get('full_name', 'Not provided') or 'Not provided',
                    ben.get('dob', 'Not provided') or 'Not provided',
                    ben.get('relationship', 'Not provided') or 'Not provided',
                    format_percentage(ben.get('allocation')),
                ])
            table = Table(table_data, hAlign='LEFT')
            table.setStyle(TableStyle([
                ('BACKGROUND', (0, 0), (-1, 0), colors.lightgrey),
                ('GRID', (0, 0), (-1, -1), 0.5, colors.grey),
            ]))
            content.append(table)
        else:
            content.append(Paragraph("[No beneficiaries specified]", normal_style))
        content.append(Spacer(1, 12))

        # Asset Breakdown
        content.append(Paragraph("Asset Breakdown", heading_style))
        asset_types = [
            "Stocks", "Bonds", "Mutual Funds", "ETFs", "UITs", 
            "Annuities (Fixed)", "Annuities (Variable)", "Options", 
            "Commodities", "Alternative Investments", "Limited Partnerships", 
            "Variable Contracts", "Short-Term", "Other"
        ]
        for asset_type in asset_types:
            field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}"
            value = form_data.get(field_name)
            content.append(Paragraph(f"{asset_type}: {format_percentage(value)}", normal_style))
        content.append(Spacer(1, 12))

        # Investment Experience
        content.append(Paragraph("Investment Experience", heading_style))
        asset_map = [
            ("Stocks", "stocks"),
            ("Bonds", "bonds"),
            ("Mutual Funds", "mutual_funds"),
            ("UITs", "uits"),
            ("Annuities (Fixed)", "annuities_fixed"),
            ("Annuities (Variable)", "annuities_variable"),
            ("Options", "options"),
            ("Commodities", "commodities"),
            ("Alternative Investments", "alternative_investments"),
            ("Limited Partnerships", "limited_partnerships"),
            ("Variable Contracts", "variable_contracts"),
        ]
        for label, key in asset_map:
            year = form_data.get(f"{key}_year_started") or "[Not provided]"
            level = form_data.get(f"{key}_level") or "[Not provided]"
            content.append(
                Paragraph(
                    f"{label} – Year Started: {year}, Level: {level}", normal_style
                )
            )
        content.append(Spacer(1, 12))

        # Broker-Dealer Relationships
        bd_primary = [
            form_data.get("employee_this_bd", "No"),
            form_data.get("related_this_bd", "No"),
            form_data.get("employee_other_bd", "No"),
            form_data.get("related_other_bd", "No"),
        ]
        bd_rows = []
        if form_data.get("employee_this_bd") == "Yes":
            bd_rows.extend([
                ("Employee Name", form_data.get("employee_name")),
                ("Department", form_data.get("department")),
                ("Branch", form_data.get("branch")),
                ("Start Date", form_data.get("start_date")),
            ])
        if form_data.get("related_this_bd") == "Yes":
            bd_rows.extend([
                ("Related Employee Name", form_data.get("related_employee_name")),
                ("Relationship", form_data.get("relationship")),
                ("Branch", form_data.get("branch_related")),
            ])
        if form_data.get("employee_other_bd") == "Yes":
            bd_rows.extend([
                ("Firm Name", form_data.get("firm_name_other")),
                ("CRD", form_data.get("crd_other")),
                ("Role", form_data.get("role_other")),
                ("Start Date", form_data.get("start_date_other")),
            ])
        if form_data.get("related_other_bd") == "Yes":
            bd_rows.extend([
                ("Firm Name", form_data.get("firm_name_rel")),
                ("Employee Name", form_data.get("employee_name_rel")),
                ("Relationship", form_data.get("relationship_rel")),
            ])
        render_section(
            "Broker-Dealer Relationships",
            bd_rows if "Yes" in bd_primary else None,
        )

        # Regulatory Affiliations
        reg_primary = [
            form_data.get("sro_member", "No"),
            form_data.get("control_person", "No"),
        ]
        reg_rows = []
        if form_data.get("sro_member") == "Yes":
            reg_rows.extend([
                ("Membership Type", form_data.get("membership_type")),
                ("CRD/Member ID", form_data.get("sro_crd")),
                ("Branch", form_data.get("sro_branch")),
            ])
        if form_data.get("control_person") == "Yes":
            reg_rows.extend([
                ("Company", form_data.get("company_name")),
                ("Ticker", form_data.get("ticker")),
                ("Exchange", form_data.get("exchange")),
                ("Role/Capacity", form_data.get("role")),
                ("Ownership %", format_pct(form_data.get("ownership_pct"))),
                ("As Of", form_data.get("as_of")),
            ])
        render_section(
            "Regulatory Affiliations",
            reg_rows if "Yes" in reg_primary else None,
        )

        # Foreign Financial Accounts
        foreign_primary = [form_data.get("has_ffi", "No")]
        foreign_rows = []
        if form_data.get("has_ffi") == "Yes":
            foreign_rows.extend([
                ("Institution", form_data.get("institution_name")),
                ("Country", form_data.get("country")),
                ("Purpose", form_data.get("purpose")),
                ("Source of Funds", form_data.get("source_of_funds")),
                ("Open Date", form_data.get("open_date")),
                (
                    "Private Banking Account",
                    "Yes" if form_data.get("private_banking") in (True, "Yes") else "No",
                ),
                (
                    "Foreign Bank Account",
                    "Yes" if form_data.get("foreign_bank_acct") in (True, "Yes") else "No",
                ),
            ])
        render_section(
            "Foreign Financial Accounts",
            foreign_rows if "Yes" in foreign_primary else None,
        )

        # Politically Exposed Person (PEP)
        pep_primary = [form_data.get("is_pep", "No")]
        pep_rows = []
        if form_data.get("is_pep") == "Yes":
            pep_rows.extend([
                ("Name", form_data.get("pep_name")),
                ("Country", form_data.get("pep_country")),
                ("Relationship", form_data.get("pep_relationship")),
                ("Title", form_data.get("pep_title")),
                ("Start Date", form_data.get("pep_start")),
                ("End Date", form_data.get("pep_end")),
                (
                    "Screening Consent",
                    "Yes" if form_data.get("pep_screening_consent") in (True, "Yes") else "No",
                ),
            ])
        render_section(
            "Politically Exposed Person (PEP)",
            pep_rows if "Yes" in pep_primary else None,
        )

        # Outside Broker Information
        if form_data.get('outside_broker_assets'):
            content.append(Paragraph("Outside Broker Information", heading_style))
            content.append(Paragraph(f"Broker Firm Name: {form_data.get('outside_firm_name', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Account Type: {form_data.get('outside_broker_account_type', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Account Number: {form_data.get('outside_broker_account_number', '[Not provided]')}", normal_style))
            content.append(Paragraph(f"Liquid Amount: {format_money(form_data.get('outside_liquid_amount'))}", normal_style))
            content.append(Spacer(1, 12))

        # Trusted Contact Information
        content.append(Paragraph("Trusted Contact Information", heading_style))
        content.append(Paragraph(f"Full Name: {form_data.get('trusted_full_name', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Relationship: {form_data.get('trusted_relationship', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Phone Number: {form_data.get('trusted_phone', '[Not provided]')}", normal_style))
        content.append(Paragraph(f"Email Address: {form_data.get('trusted_email', '[Not provided]')}", normal_style))
        content.append(Spacer(1, 12))

        # Regulatory Consent
        content.append(Paragraph("Regulatory Consent", heading_style))
        electronic_consent = form_data.get('ed_consent', 'No') or 'No'
        content.append(Paragraph(f"Electronic Delivery Consent: {electronic_consent}", normal_style))
        
        # Add page numbers
        def add_page_number(canvas, doc):
            canvas.saveState()
            canvas.setFont('Helvetica', 8)
            page_number_text = f"Page {doc.page}"
            canvas.drawCentredString(
                doc.pagesize[0] / 2,
                0.75 * inch,
                page_number_text
            )
            canvas.restoreState()
        
        # Build the PDF with page numbers
        doc.build(content, onFirstPage=add_page_number, onLaterPages=add_page_number)
        return True
    
    except Exception as e:
        print(f"Error generating PDF: {str(e)}")
        traceback.print_exc()
        return False

# Alias for backward compatibility with main_enhanced.py
generate_pdf_from_data = generate_pdf_report


def generate(form_data, output_path):
    """Compatibility wrapper expected by the UI."""
    return generate_pdf_report(form_data, output_path)
