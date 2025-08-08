#!/usr/bin/env python3
"""
Magnus Client Intake Form - Enhanced Version 2.2
Professional client intake form for financial services with comprehensive validation,
security features, and PDF generation capabilities.
"""

import sys
import os
import json
import tempfile
from datetime import datetime
from typing import Dict, Any, List

from PyQt6.QtWidgets import (
    QApplication, QMainWindow, QWidget, QVBoxLayout, QHBoxLayout,
    QLabel, QLineEdit, QPushButton, QStackedWidget, QFrame,
    QComboBox, QDateEdit, QTextEdit, QCheckBox, QRadioButton,
    QButtonGroup, QSpinBox, QGroupBox, QScrollArea, QMessageBox,
    QProgressBar, QFileDialog,
)
import traceback
from PyQt6.QtCore import Qt, QDate, QTimer, pyqtSignal
from PyQt6.QtGui import QFont, QPixmap, QIcon
import docx
import subprocess
import platform
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches

# Import custom modules
from security import DataSecurity
from pdf_generator_reportlab import generate_pdf_from_data
from ui.widgets.enhanced_line_edit import EnhancedLineEdit
from ui.pages import PAGE_CREATORS


class MagnusClientIntakeForm(QMainWindow):
    """Main application window for Magnus Client Intake Form"""
    
    def __init__(self):
        super().__init__()
        self.security_manager = DataSecurity()
        self.form_data = {}
        self.current_page = 0
        self.auto_save_timer = QTimer()
        self.auto_save_timer.timeout.connect(self.auto_save_data)
        self.auto_save_timer.start(30000)  # Auto-save every 30 seconds
        
        self.init_ui()
        self.load_draft_data()
        
    def init_ui(self):
        """Initialize the user interface"""
        self.setWindowTitle("Magnus Client Intake Form v2.2")
        self.setGeometry(100, 100, 900, 700)
        
        # Set application icon
        if os.path.exists("ICON.ico"):
            self.setWindowIcon(QIcon("ICON.ico"))
        
        # Create central widget and main layout
        central_widget = QWidget()
        self.setCentralWidget(central_widget)
        
        main_layout = QVBoxLayout(central_widget)
        main_layout.setContentsMargins(20, 20, 20, 20)
        
        # Header
        header_layout = QHBoxLayout()
        
        # Title
        title_label = QLabel("Magnus Client Intake Form")
        title_font = QFont()
        title_font.setPointSize(18)
        title_font.setBold(True)
        title_label.setFont(title_font)
        title_label.setStyleSheet("color: #2c3e50; margin-bottom: 10px;")
        header_layout.addWidget(title_label)
        
        # Progress bar
        self.progress_bar = QProgressBar()
        self.progress_bar.setValue(1)
        self.progress_bar.setStyleSheet("""
            QProgressBar {
                border: 2px solid #bdc3c7;
                border-radius: 5px;
                text-align: center;
                font-weight: bold;
            }
            QProgressBar::chunk {
                background-color: #3498db;
                border-radius: 3px;
            }
        """)
        header_layout.addWidget(self.progress_bar)

        main_layout.addLayout(header_layout)

        # Create stacked widget for pages
        self.stacked_widget = QStackedWidget()
        main_layout.addWidget(self.stacked_widget)

        # Dynamically create pages
        for creator in PAGE_CREATORS:
            creator(self)

        # Track total pages
        self.num_pages = len(PAGE_CREATORS)
        self.progress_bar.setMaximum(self.num_pages)

        # Status bar
        self.statusBar().showMessage(f"Ready - Page 1 of {self.num_pages}")
        
    def create_navigation_buttons(self, back_index=None, next_index=None):
        """Create navigation buttons layout"""
        layout = QHBoxLayout()
        
        if back_index is not None:
            back_btn = QPushButton("← Back")
            back_btn.clicked.connect(lambda: self.navigate_to_page(back_index))
            back_btn.setStyleSheet("""
                QPushButton {
                    background-color: #6c757d;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 5px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #5a6268;
                }
            """)
            layout.addWidget(back_btn)
        
        layout.addStretch()
        
        if next_index is not None:
            next_btn = QPushButton("Next →")
            next_btn.clicked.connect(lambda: self.navigate_to_page(next_index))
            next_btn.setStyleSheet("""
                QPushButton {
                    background-color: #007bff;
                    color: white;
                    border: none;
                    padding: 10px 20px;
                    border-radius: 5px;
                    font-weight: bold;
                }
                QPushButton:hover {
                    background-color: #0056b3;
                }
            """)
            layout.addWidget(next_btn)
        
        return layout

    def add_dependent_field(self, dependent_data=None):
        """Add fields for a new dependent"""
        dependent_frame = QFrame()
        dependent_frame.setFrameShape(QFrame.Shape.StyledPanel)
        dependent_frame.setFrameShadow(QFrame.Shadow.Raised)
        dependent_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                background-color: #f0f0f0;
            }
        """)

        frame_layout = QVBoxLayout(dependent_frame)

        frame_layout.addWidget(QLabel("Dependent Full Name:"))
        name_input = EnhancedLineEdit("dependent_name")
        name_input.setObjectName(f"dependent_name_{self.dependents_layout.count()}")
        frame_layout.addWidget(name_input)

        frame_layout.addWidget(QLabel("Dependent Date of Birth:"))
        dob_input = QDateEdit()
        dob_input.setObjectName(f"dependent_dob_{self.dependents_layout.count()}")
        dob_input.setDate(QDate.currentDate().addYears(-10))
        dob_input.setCalendarPopup(True)
        dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(dob_input)

        frame_layout.addWidget(QLabel("Relationship:"))
        relationship_input = EnhancedLineEdit("dependent_relationship")
        relationship_input.setObjectName(f"dependent_relationship_{self.dependents_layout.count()}")
        frame_layout.addWidget(relationship_input)

        remove_btn = QPushButton("Remove Dependent")
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        remove_btn.clicked.connect(lambda: self.remove_dependent_field(dependent_frame))
        frame_layout.addWidget(remove_btn)

        self.dependents_layout.addWidget(dependent_frame)

        if dependent_data:
            name_input.setText(dependent_data.get("name", ""))
            dob_input.setDate(QDate.fromString(dependent_data.get("dob", ""), "MM/dd/yyyy"))
            relationship_input.setText(dependent_data.get("relationship", ""))

    def remove_dependent_field(self, frame):
        """Remove dependent fields"""
        frame.deleteLater()

    def add_beneficiary_field(self, beneficiary_data=None):
        """Add fields for a new beneficiary"""
        beneficiary_frame = QFrame()
        beneficiary_frame.setFrameShape(QFrame.Shape.StyledPanel)
        beneficiary_frame.setFrameShadow(QFrame.Shadow.Raised)
        beneficiary_frame.setStyleSheet("""
            QFrame {
                border: 1px solid #ccc;
                border-radius: 5px;
                padding: 10px;
                background-color: #f0f0f0;
            }
        """)

        frame_layout = QVBoxLayout(beneficiary_frame)

        frame_layout.addWidget(QLabel("Beneficiary Full Name:"))
        name_input = EnhancedLineEdit("beneficiary_name")
        name_input.setObjectName(f"beneficiary_name_{self.beneficiaries_layout.count()}")
        frame_layout.addWidget(name_input)

        frame_layout.addWidget(QLabel("Beneficiary Date of Birth:"))
        dob_input = QDateEdit()
        dob_input.setObjectName(f"beneficiary_dob_{self.beneficiaries_layout.count()}")
        dob_input.setDate(QDate.currentDate().addYears(-10))
        dob_input.setCalendarPopup(True)
        dob_input.setStyleSheet("""
            QDateEdit {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(dob_input)

        frame_layout.addWidget(QLabel("Relationship:"))
        relationship_input = EnhancedLineEdit("beneficiary_relationship")
        relationship_input.setObjectName(f"beneficiary_relationship_{self.beneficiaries_layout.count()}")
        frame_layout.addWidget(relationship_input)

        frame_layout.addWidget(QLabel("Allocation Percentage (%):"))
        percentage_spin = QSpinBox()
        percentage_spin.setObjectName(f"beneficiary_percentage_{self.beneficiaries_layout.count()}")
        percentage_spin.setRange(0, 100)
        percentage_spin.setSuffix("%")
        percentage_spin.setStyleSheet("""
            QSpinBox {
                padding: 8px;
                border: 2px solid #ddd;
                border-radius: 4px;
                font-size: 12px;
            }
        """)
        frame_layout.addWidget(percentage_spin)

        remove_btn = QPushButton("Remove Beneficiary")
        remove_btn.setStyleSheet("""
            QPushButton {
                background-color: #dc3545;
                color: white;
                border: none;
                padding: 5px 10px;
                border-radius: 3px;
                font-weight: bold;
            }
            QPushButton:hover {
                background-color: #c82333;
            }
        """)
        remove_btn.clicked.connect(lambda: self.remove_beneficiary_field(beneficiary_frame))
        frame_layout.addWidget(remove_btn)

        self.beneficiaries_layout.addWidget(beneficiary_frame)

        if beneficiary_data:
            name_input.setText(beneficiary_data.get("name", ""))
            dob_input.setDate(QDate.fromString(beneficiary_data.get("dob", ""), "MM/dd/yyyy"))
            relationship_input.setText(beneficiary_data.get("relationship", ""))
            percentage_spin.setValue(beneficiary_data.get("percentage", 0))

    def remove_beneficiary_field(self, frame):
        """Remove beneficiary fields"""
        frame.deleteLater()

    def on_include_breakdown_changed(self, state):
        """Handle include breakdown checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        self.asset_breakdown_group.setVisible(is_checked)

    def on_has_outside_broker_changed(self, state):
        """Handle has outside broker checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        self.outside_broker_group.setVisible(is_checked)

    def on_employment_status_changed(self, text):
        """Handle employment status change to show/hide retirement fields"""
        if text == "Retired":
            self.retirement_group.setVisible(True)
        else:
            self.retirement_group.setVisible(False)

    def on_spouse_applicable_changed(self, state):
        """Handle spouse applicable checkbox change"""
        is_checked = state == Qt.CheckState.Checked.value
        self.spouse_group.setVisible(is_checked)
        
    def navigate_to_page(self, page_index):
        """Navigate to a specific page"""
        if page_index == self.num_pages - 1:  # Review page
            self.update_review_area()

        self.current_page = page_index
        self.stacked_widget.setCurrentIndex(page_index)
        self.progress_bar.setValue(page_index + 1)
        self.statusBar().showMessage(f"Page {page_index + 1} of {self.num_pages}")
        
        # Auto-save when navigating
        self.collect_form_data()
        
    def update_review_area(self):
        """Update the review area with current form data"""
        self.collect_form_data()
        
        review_text = "=== MAGNUS CLIENT INTAKE FORM - REVIEW ===\n\n"
        
        # Helper to format fields
        def format_field(label, value):
            return f"  {label}: {value if value else '[Not provided]'}\n"

        # Personal Information
        review_text += "PERSONAL INFORMATION:\n"
        review_text += format_field("Full Name", self.form_data.get("full_name"))
        review_text += format_field("Date of Birth", self.form_data.get("dob"))
        review_text += format_field("Social Security Number", self.form_data.get("ssn"))
        review_text += format_field("Citizenship", self.form_data.get("citizenship"))
        review_text += format_field("Marital Status", self.form_data.get("marital_status"))
        review_text += "\n"

        # Contact Information
        review_text += "CONTACT INFORMATION:\n"
        review_text += format_field("Residential Address", self.form_data.get("residential_address"))
        if self.form_data.get("mailing_address_different"):
            review_text += format_field("Mailing Address", self.form_data.get("mailing_address"))
        review_text += format_field("Email", self.form_data.get("email"))
        review_text += format_field("Home Phone", self.form_data.get("home_phone"))
        review_text += format_field("Mobile Phone", self.form_data.get("mobile_phone"))
        review_text += format_field("Work Phone", self.form_data.get("work_phone"))
        review_text += "\n"

        # Employment Information
        review_text += "EMPLOYMENT INFORMATION:\n"
        review_text += format_field("Employment Status", self.form_data.get("employment_status"))
        review_text += format_field("Employer Name", self.form_data.get("employer_name"))
        review_text += format_field("Occupation", self.form_data.get("occupation"))
        review_text += format_field("Years Employed", self.form_data.get("years_employed"))
        review_text += format_field("Annual Income", self.form_data.get("annual_income"))
        review_text += format_field("Employer Address", self.form_data.get("employer_address"))
        review_text += "\n"

        # Retirement Information
        if self.form_data.get("employment_status") == "Retired":
            review_text += "RETIREMENT INFORMATION:\n"
            review_text += format_field("Former Employer", self.form_data.get("former_employer"))
            review_text += format_field("Source of Income", self.form_data.get("income_source"))
            review_text += "\n"

        # Financial Information
        review_text += "FINANCIAL INFORMATION:\n"
        review_text += format_field("Education Status", self.form_data.get("education_status"))
        review_text += format_field("Estimated Tax Bracket", self.form_data.get("tax_bracket"))
        review_text += format_field("Investment Risk Tolerance", self.form_data.get("risk_tolerance"))
        review_text += format_field("Investment Purpose", self.form_data.get("investment_purpose"))
        review_text += format_field("Investment Objectives", self.form_data.get("investment_objective"))
        review_text += format_field("Net Worth (excluding primary home)", self.form_data.get("net_worth"))
        review_text += format_field("Liquid Net Worth", self.form_data.get("liquid_net_worth"))
        review_text += format_field("Assets Held Away", self.form_data.get("assets_held_away"))
        review_text += "\n"

        # Spouse Information
        if not self.form_data.get("spouse_applicable"):
            review_text += "SPOUSE INFORMATION:\n"
            review_text += format_field("Spouse Full Name", self.form_data.get("spouse_full_name"))
            review_text += format_field("Spouse Date of Birth", self.form_data.get("spouse_dob"))
            review_text += format_field("Spouse SSN", self.form_data.get("spouse_ssn"))
            review_text += format_field("Spouse Employment Status", self.form_data.get("spouse_employment_status"))
            review_text += format_field("Spouse Employer Name", self.form_data.get("spouse_employer_name"))
            review_text += format_field("Spouse Occupation/Title", self.form_data.get("spouse_occupation"))
            review_text += "\n"
        else:
            review_text += "SPOUSE INFORMATION:\n  [Not applicable]\n\n"

        # Dependents
        review_text += "DEPENDENTS:\n"
        dependents = self.form_data.get("dependents", [])
        if dependents:
            for i, dep in enumerate(dependents):
                review_text += f"  Dependent {i+1}:\n"
                review_text += format_field("    Name", dep.get("name"))
                review_text += format_field("    Date of Birth", dep.get("dob"))
                review_text += format_field("    Relationship", dep.get("relationship"))
        else:
            review_text += "  [No dependents specified]\n"
        review_text += "\n"

        # Beneficiaries
        review_text += "BENEFICIARIES:\n"
        beneficiaries = self.form_data.get("beneficiaries", [])
        if beneficiaries:
            for i, ben in enumerate(beneficiaries):
                review_text += f"  Beneficiary {i+1}:\n"
                review_text += format_field("    Name", ben.get("name"))
                review_text += format_field("    Date of Birth", ben.get("dob"))
                review_text += format_field("    Relationship", ben.get("relationship"))
                percentage = ben.get('percentage', '')
                review_text += format_field("    Percentage", f"{percentage}%" if percentage else "[Not provided]")
        else:
            review_text += "  [No beneficiaries specified]\n"
        review_text += "\n"

        # Asset Breakdown
        review_text += "ASSET BREAKDOWN:\n"
        asset_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures", "Short-Term", "Other"]
        for asset_type in asset_types:
            field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
            value = self.form_data.get(field_name)
            review_text += format_field(asset_type, f"{value}%" if value else None)
        review_text += "\n"

        # Investment Experience
        review_text += "INVESTMENT EXPERIENCE:\n"
        experience_types = ["Stocks", "Bonds", "Mutual Funds", "ETFs", "Options", "Futures"]
        for exp_type in experience_types:
            year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
            level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
            
            year = self.form_data.get(year_field)
            level = self.form_data.get(level_field)
            
            review_text += f"  {exp_type}:\n"
            review_text += format_field("    Year Started", year)
            review_text += format_field("    Experience Level", level)
        review_text += "\n"

        # Outside Broker Information
        if self.form_data.get("has_outside_broker"):
            review_text += "OUTSIDE BROKER INFORMATION:\n"
            review_text += format_field("Broker Firm Name", self.form_data.get("outside_firm_name"))
            review_text += format_field("Account Number", self.form_data.get("outside_broker_account_number"))
            review_text += format_field("Account Type", self.form_data.get("outside_broker_account_type"))
            review_text += "\n"

        # Trusted Contact Information
        review_text += "TRUSTED CONTACT INFORMATION:\n"
        review_text += format_field("Full Name", self.form_data.get("trusted_full_name"))
        review_text += format_field("Relationship", self.form_data.get("trusted_relationship"))
        review_text += format_field("Phone Number", self.form_data.get("trusted_phone"))
        review_text += format_field("Email Address", self.form_data.get("trusted_email"))
        review_text += "\n"

        # Regulatory Consent
        review_text += "REGULATORY CONSENT:\n"
        electronic_consent = "Yes" if self.form_data.get("electronic_regulatory_yes") else "No"
        review_text += format_field("Electronic Delivery Consent", electronic_consent)
        review_text += "\n"
        
        self.review_area.setPlainText(review_text)
        
    def collect_form_data(self):
        """Collect all form data from the UI"""
        # Get all widgets with object names
        for widget in self.findChildren((QLineEdit, QComboBox, QDateEdit, QTextEdit, QSpinBox, QCheckBox, QRadioButton)):
            object_name = widget.objectName()
            if object_name:
                if isinstance(widget, QLineEdit):
                    self.form_data[object_name] = widget.text()
                elif isinstance(widget, QComboBox):
                    self.form_data[object_name] = widget.currentText()
                elif isinstance(widget, QDateEdit):
                    self.form_data[object_name] = widget.date().toString("MM/dd/yyyy")
                elif isinstance(widget, QTextEdit):
                    self.form_data[object_name] = widget.toPlainText()
                elif isinstance(widget, QSpinBox):
                    self.form_data[object_name] = widget.value()
                elif isinstance(widget, QCheckBox):
                    self.form_data[object_name] = widget.isChecked()
                elif isinstance(widget, QRadioButton):
                    if widget.isChecked():
                        self.form_data[object_name] = True

        # Collect investment purpose data
        investment_purposes = []
        for purpose, checkbox in self.purpose_checkboxes.items():
            if checkbox.isChecked():
                investment_purposes.append(purpose)
        self.form_data["investment_purpose"] = ", ".join(investment_purposes) if investment_purposes else None

        # Collect investment objectives data
        investment_objectives = []
        for objective, spinbox in self.objective_spinboxes.items():
            rank = spinbox.value()
            if rank > 0:  # Only include if a rank is selected
                investment_objectives.append(f"{objective}: {rank}")
        self.form_data["investment_objective"] = "\n".join(investment_objectives) if investment_objectives else None

        # Collect dependents data
        dependents = []
        for i in range(self.dependents_layout.count()):
            frame = self.dependents_layout.itemAt(i).widget()
            if isinstance(frame, QFrame):
                dependent_data = {}
                for child in frame.findChildren((QLineEdit, QDateEdit)):
                    if isinstance(child, QLineEdit):
                        if "name" in child.objectName():
                            dependent_data["name"] = child.text()
                        elif "relationship" in child.objectName():
                            dependent_data["relationship"] = child.text()
                    elif isinstance(child, QDateEdit):
                        dependent_data["dob"] = child.date().toString("MM/dd/yyyy")
                if dependent_data:
                    dependents.append(dependent_data)
        self.form_data["dependents"] = dependents

        # Collect beneficiaries data
        beneficiaries = []
        for i in range(self.beneficiaries_layout.count()):
            frame = self.beneficiaries_layout.itemAt(i).widget()
            if isinstance(frame, QFrame):
                beneficiary_data = {}
                for child in frame.findChildren((QLineEdit, QDateEdit, QSpinBox)):
                    if isinstance(child, QLineEdit):
                        if "name" in child.objectName():
                            beneficiary_data["name"] = child.text()
                        elif "relationship" in child.objectName():
                            beneficiary_data["relationship"] = child.text()
                    elif isinstance(child, QDateEdit):
                        beneficiary_data["dob"] = child.date().toString("MM/dd/yyyy")
                    elif isinstance(child, QSpinBox):
                        beneficiary_data["percentage"] = child.value()
                if beneficiary_data:
                    beneficiaries.append(beneficiary_data)
        self.form_data["beneficiaries"] = beneficiaries

        # Collect asset breakdown data
        if hasattr(self, 'asset_breakdown_fields'):
            for asset_type, spin_box in self.asset_breakdown_fields.items():
                field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_')}"
                self.form_data[field_name] = spin_box.value()

        # Collect investment experience data
        if hasattr(self, 'asset_experience_fields'):
            for exp_type, fields in self.asset_experience_fields.items():
                year_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_year"
                level_field = f"asset_experience_{exp_type.lower().replace(' ', '_')}_level"
                self.form_data[year_field] = fields["year"].text()
                self.form_data[level_field] = fields["level"].currentText()
                        
    def auto_save_data(self):
        """Auto-save form data"""
        # Removed auto-save functionality to prevent JSON format saving
        pass
            
    def load_draft_data(self):
        """Load draft data if available"""
        try:
            temp_file = os.path.join(tempfile.gettempdir(), "magnus_form_autosave.json")
            if os.path.exists(temp_file):
                with open(temp_file, 'r') as f:
                    self.form_data = json.load(f)
                self.populate_form_fields()
        except Exception as e:
            print(f"Failed to load draft: {e}")
            
    def populate_form_fields(self):
        """Populate form fields with loaded data"""
        for object_name, value in self.form_data.items():
            widget = self.findChild((QLineEdit, QComboBox, QDateEdit, QTextEdit, QSpinBox, QCheckBox, QRadioButton), object_name)
            if widget:
                try:
                    if isinstance(widget, QLineEdit):
                        widget.setText(str(value))
                    elif isinstance(widget, QComboBox):
                        index = widget.findText(str(value))
                        if index >= 0:
                            widget.setCurrentIndex(index)
                    elif isinstance(widget, QDateEdit):
                        widget.setDate(QDate.fromString(str(value), "MM/dd/yyyy"))
                    elif isinstance(widget, QTextEdit):
                        widget.setPlainText(str(value))
                    elif isinstance(widget, QSpinBox):
                        if value:
                            widget.setValue(int(value))
                    elif isinstance(widget, QCheckBox):
                        widget.setChecked(bool(value))
                    elif isinstance(widget, QRadioButton):
                        widget.setChecked(bool(value))
                except Exception as e:
                    print(f"Failed to populate field {object_name}: {e}")
                    
    def save_draft(self):
        """Save current form as draft in Word format"""
        self.collect_form_data()
        
        file_path, _ = QFileDialog.getSaveFileName(
            self, "Save Draft", "magnus_form_draft.docx", "Word Files (*.docx)"
        )
        
        if file_path:
            try:
                doc = Document()
                
                # Title
                title = doc.add_heading('Magnus Client Intake Form', 0)
                title.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Helper function to format monetary values
                def format_money(value):
                    if value:
                        try:
                            return f"${int(value):,}"
                        except ValueError:
                            return value
                    return "[Not provided]"
                
                # Personal Information
                doc.add_heading('Personal Information', level=1)
                doc.add_paragraph(f"Full Name: {self.form_data.get('full_name', '[Not provided]')}")
                doc.add_paragraph(f"Date of Birth: {self.form_data.get('dob', '[Not provided]')}")
                doc.add_paragraph(f"Social Security Number: {self.form_data.get('ssn', '[Not provided]')}")
                doc.add_paragraph(f"Citizenship: {self.form_data.get('citizenship', '[Not provided]')}")
                doc.add_paragraph(f"Marital Status: {self.form_data.get('marital_status', '[Not provided]')}")
                doc.add_paragraph()
                
                # Contact Information
                doc.add_heading('Contact Information', level=1)
                doc.add_paragraph(f"Residential Address: {self.form_data.get('residential_address', '[Not provided]')}")
                doc.add_paragraph(f"Email: {self.form_data.get('email', '[Not provided]')}")
                doc.add_paragraph(f"Home Phone: {self.form_data.get('home_phone', '[Not provided]')}")
                doc.add_paragraph(f"Mobile Phone: {self.form_data.get('mobile_phone', '[Not provided]')}")
                doc.add_paragraph(f"Work Phone: {self.form_data.get('work_phone', '[Not provided]')}")
                doc.add_paragraph()
                
                # Employment Information
                doc.add_heading('Employment Information', level=1)
                doc.add_paragraph(f"Employment Status: {self.form_data.get('employment_status', '[Not provided]')}")
                doc.add_paragraph(f"Employer Name: {self.form_data.get('employer_name', '[Not provided]')}")
                doc.add_paragraph(f"Occupation: {self.form_data.get('occupation', '[Not provided]')}")
                doc.add_paragraph(f"Years Employed: {self.form_data.get('years_employed', '[Not provided]')}")
                doc.add_paragraph(f"Annual Income: {format_money(self.form_data.get('annual_income'))}")
                
                # Retirement Information (if applicable)
                if self.form_data.get('employment_status') == 'Retired':
                    doc.add_heading('Retirement Information', level=1)
                    doc.add_paragraph(f"Former Employer: {self.form_data.get('former_employer', '[Not provided]')}")
                    doc.add_paragraph(f"Source of Income: {self.form_data.get('income_source', '[Not provided]')}")
                doc.add_paragraph()
                
                # Financial Information
                doc.add_heading('Financial Information', level=1)
                doc.add_paragraph(f"Education Status: {self.form_data.get('education_status', '[Not provided]')}")
                doc.add_paragraph(f"Estimated Tax Bracket: {self.form_data.get('tax_bracket', '[Not provided]')}")
                doc.add_paragraph(f"Investment Risk Tolerance: {self.form_data.get('risk_tolerance', '[Not provided]')}")
                
                # Investment Purpose
                doc.add_paragraph("Investment Purpose:")
                purpose_list = []
                for purpose in ["Income", "Growth and Income", "Capital Appreciation", "Speculation"]:
                    if self.form_data.get(f"investment_purpose_{purpose.lower().replace(' ', '_')}"):
                        purpose_list.append(purpose)
                doc.add_paragraph(", ".join(purpose_list) if purpose_list else "[Not provided]")
                
                # Investment Objectives
                doc.add_paragraph("Investment Objectives (Ranked 1-5):")
                objectives = [
                    "Trading Profits", "Speculation", "Capital Appreciation", 
                    "Income", "Preservation of Capital"
                ]
                for objective in objectives:
                    rank = self.form_data.get(f"investment_objective_{objective.lower().replace(' ', '_')}")
                    if rank:
                        doc.add_paragraph(f"  {objective}: {rank}")
                
                doc.add_paragraph(f"Net Worth: {format_money(self.form_data.get('net_worth'))}")
                doc.add_paragraph(f"Liquid Net Worth: {format_money(self.form_data.get('liquid_net_worth'))}")
                doc.add_paragraph(f"Assets Held Away: {format_money(self.form_data.get('assets_held_away'))}")
                doc.add_paragraph()
                
                # Spouse Information
                if self.form_data.get('spouse_applicable'):
                    doc.add_heading('Spouse Information', level=1)
                    doc.add_paragraph(f"Full Name: {self.form_data.get('spouse_full_name', '[Not provided]')}")
                    doc.add_paragraph(f"Date of Birth: {self.form_data.get('spouse_dob', '[Not provided]')}")
                    doc.add_paragraph(f"Social Security Number: {self.form_data.get('spouse_ssn', '[Not provided]')}")
                    doc.add_paragraph(f"Employment Status: {self.form_data.get('spouse_employment_status', '[Not provided]')}")
                    doc.add_paragraph(f"Employer Name: {self.form_data.get('spouse_employer_name', '[Not provided]')}")
                    doc.add_paragraph(f"Occupation: {self.form_data.get('spouse_occupation', '[Not provided]')}")
                    doc.add_paragraph()
                
                # Dependents
                doc.add_heading('Dependents', level=1)
                dependents = self.form_data.get('dependents', [])
                if dependents:
                    for i, dep in enumerate(dependents, 1):
                        doc.add_paragraph(f"Dependent {i}:")
                        doc.add_paragraph(f"  Name: {dep.get('name', '[Not provided]')}")
                        doc.add_paragraph(f"  Date of Birth: {dep.get('dob', '[Not provided]')}")
                        doc.add_paragraph(f"  Relationship: {dep.get('relationship', '[Not provided]')}")
                else:
                    doc.add_paragraph("[No dependents specified]")
                doc.add_paragraph()
                
                # Beneficiaries
                doc.add_heading('Beneficiaries', level=1)
                beneficiaries = self.form_data.get('beneficiaries', [])
                if beneficiaries:
                    for i, ben in enumerate(beneficiaries, 1):
                        doc.add_paragraph(f"Beneficiary {i}:")
                        doc.add_paragraph(f"  Name: {ben.get('name', '[Not provided]')}")
                        doc.add_paragraph(f"  Date of Birth: {ben.get('dob', '[Not provided]')}")
                        doc.add_paragraph(f"  Relationship: {ben.get('relationship', '[Not provided]')}")
                        percentage = ben.get('percentage', '')
                        doc.add_paragraph(f"  Percentage: {f'{percentage}%' if percentage else '[Not provided]'}")
                else:
                    doc.add_paragraph("[No beneficiaries specified]")
                doc.add_paragraph()
                
                # Asset Breakdown
                doc.add_heading('Asset Breakdown', level=1)
                asset_types = [
                    "Stocks", "Bonds", "Mutual Funds", "ETFs", "UITs", 
                    "Annuities (Fixed)", "Annuities (Variable)", "Options", 
                    "Commodities", "Alternative Investments", "Limited Partnerships", 
                    "Variable Contracts", "Short-Term", "Other"
                ]
                for asset_type in asset_types:
                    field_name = f"asset_breakdown_{asset_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}"
                    value = self.form_data.get(field_name)
                    doc.add_paragraph(f"{asset_type}: {f'{value}%' if value else '[Not provided]'}")
                doc.add_paragraph()
                
                # Investment Experience
                doc.add_heading('Investment Experience', level=1)
                experience_types = [
                    "Stocks", "Bonds", "Mutual Funds", "UITs", 
                    "Annuities (Fixed)", "Annuities (Variable)", "Options", 
                    "Commodities", "Alternative Investments", "Limited Partnerships", 
                    "Variable Contracts"
                ]
                for exp_type in experience_types:
                    doc.add_paragraph(f"{exp_type}:")
                    year_field = f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_year"
                    level_field = f"asset_experience_{exp_type.lower().replace(' ', '_').replace('(', '').replace(')', '')}_level"
                    year = self.form_data.get(year_field)
                    level = self.form_data.get(level_field)
                    doc.add_paragraph(f"  Year Started: {year or '[Not provided]'}")
                    doc.add_paragraph(f"  Experience Level: {level or '[Not provided]'}")
                doc.add_paragraph()
                
                # Outside Broker Information
                if self.form_data.get('has_outside_broker'):
                    doc.add_heading('Outside Broker Information', level=1)
                    doc.add_paragraph(f"Broker Firm Name: {self.form_data.get('outside_firm_name', '[Not provided]')}")
                    doc.add_paragraph(f"Account Type: {self.form_data.get('outside_broker_account_type', '[Not provided]')}")
                    doc.add_paragraph(f"Account Number: {self.form_data.get('outside_broker_account_number', '[Not provided]')}")
                    doc.add_paragraph(f"Liquid Amount: {format_money(self.form_data.get('outside_liquid_amount'))}")
                    doc.add_paragraph()
                
                # Trusted Contact Information
                doc.add_heading('Trusted Contact Information', level=1)
                doc.add_paragraph(f"Full Name: {self.form_data.get('trusted_full_name', '[Not provided]')}")
                doc.add_paragraph(f"Relationship: {self.form_data.get('trusted_relationship', '[Not provided]')}")
                doc.add_paragraph(f"Phone Number: {self.form_data.get('trusted_phone', '[Not provided]')}")
                doc.add_paragraph(f"Email Address: {self.form_data.get('trusted_email', '[Not provided]')}")
                doc.add_paragraph()
                
                # Regulatory Consent
                doc.add_heading('Regulatory Consent', level=1)
                electronic_consent = "Yes" if self.form_data.get('electronic_regulatory_yes') else "No"
                doc.add_paragraph(f"Electronic Delivery Consent: {electronic_consent}")
                
                # Save the document
                doc.save(file_path)
                QMessageBox.information(self, "Success", "Draft saved successfully in Word format!")
            except Exception as e:
                QMessageBox.critical(self, "Error", f"Failed to save draft: {e}")
                
    def generate_pdf_report(self):
        """Generate a PDF report from the form data"""
        try:
            # First collect all form data
            self.collect_form_data()

            # Get the save location from user
            file_path, _ = QFileDialog.getSaveFileName(
                self, "Save PDF Report", "magnus_form_report.pdf", "PDF Files (*.pdf)"
            )
            
            if not file_path:  # User cancelled
                return
            
            # Generate the PDF
            from pdf_generator_reportlab import generate_pdf_report
            success = generate_pdf_report(self.form_data, file_path)
            
            if success:
                QMessageBox.information(
                    self,
                    "Success",
                    f"PDF has been generated successfully and saved to:\n{file_path}"
                )
                # Try to open the PDF
                try:
                    if os.name == 'nt':  # Windows
                        os.startfile(file_path)
                    elif os.name == 'posix':  # macOS or Linux
                        subprocess.run(['open' if sys.platform == 'darwin' else 'xdg-open', file_path])
                except Exception as e:
                    QMessageBox.warning(
                        self,
                        "Warning",
                        f"PDF was generated but could not be opened automatically:\n{str(e)}"
                    )
            else:
                QMessageBox.critical(
                    self,
                    "Error",
                    "Failed to generate PDF. Please try again."
                )
        except Exception as e:
            error_msg = f"An error occurred while generating the PDF:\n{str(e)}\n\nTraceback:\n{traceback.format_exc()}"
            print(error_msg)  # Print to console for debugging
            QMessageBox.critical(
                self,
                "Error Generating PDF",
                error_msg
            )



def main():
    """Main application entry point"""
    app = QApplication(sys.argv)
    app.setApplicationName("Magnus Client Intake Form")
    app.setApplicationVersion("2.2")
    
    # Set application style
    app.setStyleSheet("""
        QMainWindow {
            background-color: #ffffff;
        }
        QLabel {
            color: #2c3e50;
        }
        QGroupBox {
            font-weight: bold;
            border: 2px solid #bdc3c7;
            border-radius: 5px;
            margin-top: 10px;
            padding-top: 10px;
        }
        QGroupBox::title {
            subcontrol-origin: margin;
            left: 10px;
            padding: 0 5px 0 5px;
        }
    """)
    
    # Create and show main window
    window = MagnusClientIntakeForm()
    window.show()
    
    return app.exec()


if __name__ == "__main__":
    sys.exit(main())


